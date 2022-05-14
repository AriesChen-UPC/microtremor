# encoding: UTF-8
"""
@author: AriesChen
@contact: s15010125@s.upc.edu.cn
@time: 4/29/2022 3:56 PM
@file: fig_doc.py
       This script is used to generate the doc of figures for report.
"""

import shutil
from docx import Document
from docx.shared import Inches
import os
import tkinter as tk
from tkinter import filedialog
from glob import glob
import colorama
colorama.init(autoreset=True)

#%% Grabbing images

print('\033[0;31mPlease input How many folders to browse:\033[0m')
num_folders = int(input())
while num_folders > 0:
    num_folders -= 1
    print('\033[0;31mPlease select the figure file folder：\033[0m')
    root = tk.Tk()
    root.withdraw()
    doc_fig_path = filedialog.askdirectory()
    print('\033[0;31mFolderPath\033[0m:', doc_fig_path)
    fig_save_path = os.path.abspath(os.path.dirname(doc_fig_path)) + '/Attachment'
    hv_fig_save_path = fig_save_path + '/HVSR'
    spac_fig_save_path = fig_save_path + '/SPAC'
    if not os.path.exists(fig_save_path):
        os.mkdir(fig_save_path)
    if not os.path.exists(hv_fig_save_path):
        os.makedirs(hv_fig_save_path)
    if not os.path.exists(spac_fig_save_path):
        os.makedirs(spac_fig_save_path)
    sub_path_name = os.listdir(doc_fig_path)
    sub_path = []
    for i in range(len(sub_path_name)):
        sub_path.append(doc_fig_path + '/' + sub_path_name[i])
        try:
            hv_fig = glob(sub_path[i] + '/*hv*.png')
            if len(hv_fig) > 1:
                shutil.copy(hv_fig[0], hv_fig_save_path)
            else:
                shutil.copy(hv_fig[0], hv_fig_save_path)
        except IndexError:
            pass
        try:
            spac_fig = glob(sub_path[i] + '/*spacfig*.png')
            if len(spac_fig) > 1:
                shutil.copy(spac_fig[0], spac_fig_save_path)
            else:
                shutil.copy(spac_fig[0], spac_fig_save_path)
        except IndexError:
            pass

#%% Creating Word Document

hv_fig_path = glob(hv_fig_save_path + '/*.png')
doc = Document()
for fig in hv_fig_path:
    doc.add_picture(fig, width=Inches(6), height=Inches(2.5))  # todo: to align the figures in center
    paragraph = doc.add_paragraph(os.path.basename(fig).split('.')[0].split('-')[-1])
    paragraph.alignment = 1
doc.save(hv_fig_save_path + '/hv_doc.docx')
print('\033[0;32mHV figure doc created successfully!\033[0m')

spac_fig_path = glob(spac_fig_save_path + '/*.png')
doc = Document()
for fig in spac_fig_path:
    doc.add_picture(fig, width=Inches(6), height=Inches(2.5))
    paragraph = doc.add_paragraph(os.path.basename(fig).split('.')[0].split('-')[-1])
    paragraph.alignment = 1
doc.save(spac_fig_save_path + '/spac_doc.docx')
print('\033[0;32mSPAC figure doc created successfully!\033[0m')

#%% main

if __name__ == "__main__":
    pass
